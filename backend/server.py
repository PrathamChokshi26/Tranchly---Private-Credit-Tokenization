from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException, Form
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
import json
import io
import PyPDF2
import openpyxl
from docx import Document
from emergentintegrations.llm.chat import LlmChat, UserMessage

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# LLM API Key
LLM_API_KEY = os.environ.get('EMERGENT_LLM_KEY')

# ============= MODELS =============

class FinancialDocument(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str = "default_user"
    filename: str
    file_type: str
    content: str
    upload_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    metadata: Dict[str, Any] = {}

class Analysis(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_id: Optional[str] = None
    analysis_type: str
    content: str
    result: Dict[str, Any]
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Portfolio(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str = "default_user"
    name: str
    holdings: List[Dict[str, Any]]
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Company(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    ticker: Optional[str] = None
    sector: Optional[str] = None
    industry: Optional[str] = None
    tracked: bool = False
    notes: str = ""
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class AnalysisRequest(BaseModel):
    content: str
    analysis_type: str
    document_id: Optional[str] = None

class SimulationRequest(BaseModel):
    company_name: str
    revenue_growth: float
    gross_margin: float
    operating_leverage: float
    capex_intensity: float

class PortfolioCreate(BaseModel):
    name: str
    holdings: List[Dict[str, Any]]

class CompanyCreate(BaseModel):
    name: str
    ticker: Optional[str] = None
    sector: Optional[str] = None
    industry: Optional[str] = None
    notes: str = ""

# ============= HELPER FUNCTIONS =============

async def get_llm_analysis(prompt: str, context: str = "") -> str:
    """Get AI analysis using emergentintegrations"""
    try:
        chat = LlmChat(
            api_key=LLM_API_KEY,
            session_id=str(uuid.uuid4()),
            system_message="You are a financial analysis expert with deep knowledge of financial statements, accounting, valuation, and investment analysis. Provide clear, insightful analysis."
        ).with_model("openai", "gpt-5")
        
        full_prompt = f"{context}\n\n{prompt}" if context else prompt
        message = UserMessage(text=full_prompt)
        response = await chat.send_message(message)
        return response
    except Exception as e:
        logging.error(f"LLM analysis error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(e)}")

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF - Optimized"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text = ""
        # Limit to first 100 pages for speed
        max_pages = min(len(pdf_reader.pages), 100)
        for i in range(max_pages):
            page_text = pdf_reader.pages[i].extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"PDF parsing error: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"DOCX parsing error: {str(e)}")

def extract_text_from_excel(file_bytes: bytes) -> str:
    """Extract text from Excel"""
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes))
        text = ""
        for sheet in wb.worksheets:
            text += f"\n=== Sheet: {sheet.title} ===\n"
            for row in sheet.iter_rows(values_only=True):
                text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Excel parsing error: {str(e)}")

# ============= ROUTES =============

@api_router.get("/")
async def root():
    return {"message": "Financial Research AI Platform API", "status": "online"}

@api_router.get("/health")
async def health_check():
    """Quick health check endpoint"""
    try:
        # Test DB connection
        await db.command("ping")
        return {"status": "healthy", "database": "connected", "timestamp": datetime.now(timezone.utc).isoformat()}
    except Exception as e:
        return {"status": "unhealthy", "error": str(e)}

@api_router.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload and process financial document - FAST MODE"""
    try:
        start_time = datetime.now(timezone.utc)
        logger.info(f"[UPLOAD START] File: {file.filename}")
        
        # Read file with size limit (50MB max)
        max_size = 50 * 1024 * 1024  # 50MB
        file_bytes = await file.read()
        
        if len(file_bytes) > max_size:
            raise HTTPException(status_code=400, detail="File too large. Maximum size is 50MB")
        
        filename = file.filename
        file_ext = filename.split('.')[-1].lower()
        
        logger.info(f"[UPLOAD] Processing {filename}, size: {len(file_bytes)} bytes")
        
        # Quick extraction with timeout protection
        content = ""
        try:
            if file_ext == 'pdf':
                content = extract_text_from_pdf(file_bytes)
                file_type = 'pdf'
            elif file_ext in ['xlsx', 'xls']:
                content = extract_text_from_excel(file_bytes)
                file_type = 'excel'
            elif file_ext in ['docx', 'doc']:
                content = extract_text_from_docx(file_bytes)
                file_type = 'docx'
            elif file_ext == 'txt':
                content = file_bytes.decode('utf-8')
                file_type = 'txt'
            elif file_ext == 'csv':
                content = file_bytes.decode('utf-8')
                file_type = 'csv'
            else:
                raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_ext}. Supported: PDF, Excel, Word, Text, CSV")
        except Exception as parse_error:
            logger.error(f"[UPLOAD ERROR] Parse failed: {str(parse_error)}")
            raise HTTPException(status_code=400, detail=f"Failed to parse file: {str(parse_error)}")
        
        if not content or len(content) < 10:
            raise HTTPException(status_code=400, detail="No content extracted from file. File may be empty or corrupted.")
        
        logger.info(f"[UPLOAD] Extracted {len(content)} characters")
        
        # Quick DB save
        doc = FinancialDocument(
            filename=filename,
            file_type=file_type,
            content=content[:100000]  # Limit to 100k chars for storage
        )
        
        doc_dict = doc.model_dump()
        doc_dict['upload_date'] = doc_dict['upload_date'].isoformat()
        
        try:
            await db.documents.insert_one(doc_dict)
            logger.info(f"[UPLOAD SUCCESS] Saved to DB: {doc.id}")
        except Exception as db_error:
            logger.error(f"[UPLOAD ERROR] DB save failed: {str(db_error)}")
            raise HTTPException(status_code=500, detail="Failed to save document to database")
        
        elapsed = (datetime.now(timezone.utc) - start_time).total_seconds()
        logger.info(f"[UPLOAD COMPLETE] {filename} in {elapsed:.2f}s")
        
        return {
            "success": True,
            "document_id": doc.id,
            "filename": filename,
            "content_length": len(content),
            "processing_time": f"{elapsed:.2f}s",
            "message": "Document uploaded successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[UPLOAD ERROR] Unexpected: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@api_router.post("/analyze/statement")
async def analyze_statement(request: AnalysisRequest):
    """Analyze financial statement with deep insights for 10K, 10Q, Annual Reports"""
    
    # Detect document type for specialized prompts
    content_lower = request.content.lower()
    is_10k = '10-k' in content_lower or 'form 10-k' in content_lower
    is_10q = '10-q' in content_lower or 'form 10-q' in content_lower
    is_annual = 'annual report' in content_lower
    
    doc_type = "10-K filing" if is_10k else "10-Q filing" if is_10q else "Annual Report" if is_annual else "Financial Statement"
    
    prompt = f"""You are analyzing a {doc_type}. Provide COMPREHENSIVE financial analysis:

DOCUMENT CONTENT:
{request.content[:15000]}  

REQUIRED ANALYSIS (respond in valid JSON format):

{{
  "summary": "Executive summary highlighting key findings, trends, and critical insights (5-7 sentences)",
  
  "line_items": {{
    "Revenue": "Value with YoY/QoQ comparison",
    "Gross Profit": "Value with margin %",
    "Operating Income": "Value with margin %",
    "Net Income": "Value with margin %",
    "EPS": "Value with growth %",
    "Total Assets": "Value",
    "Total Liabilities": "Value",
    "Shareholders Equity": "Value",
    "Cash and Equivalents": "Value",
    "Free Cash Flow": "Value if available"
  }},
  
  "trends": {{
    "revenue": [
      {{"period": "Q1 2023", "value": 550}},
      {{"period": "Q2 2023", "value": 580}},
      {{"period": "Q3 2023", "value": 610}},
      {{"period": "Q4 2023", "value": 760}}
    ],
    "growth": [
      {{"metric": "Revenue Growth", "value": 15}},
      {{"metric": "Gross Margin", "value": 48}},
      {{"metric": "Operating Margin", "value": 18}},
      {{"metric": "Net Margin", "value": 13}},
      {{"metric": "ROE", "value": 22}}
    ]
  }},
  
  "margins": {{
    "data": [
      {{"name": "Gross Margin", "current": 48, "previous": 45.7, "industry": 42}},
      {{"name": "Operating Margin", "current": 18, "previous": 16.5, "industry": 15}},
      {{"name": "Net Margin", "current": 13, "previous": 11.8, "industry": 10}}
    ]
  }},
  
  "financial_ratios": [
    {{"category": "Liquidity", "score": 85}},
    {{"category": "Profitability", "score": 90}},
    {{"category": "Efficiency", "score": 78}},
    {{"category": "Leverage", "score": 75}},
    {{"category": "Growth", "score": 88}}
  ],
  
  "red_flags": "List any concerning trends: working capital issues, margin compression, cash flow problems, accounting irregularities, or state 'No significant red flags detected'",
  
  "health_score": 82,
  
  "key_insights": [
    "Strategic initiative or market position insight",
    "Operational efficiency observation",
    "Capital allocation or liquidity insight",
    "Risk factor or competitive pressure"
  ],
  
  "segment_analysis": "Analysis of business segments if mentioned in document",
  
  "management_discussion": "Key points from MD&A section if available"
}}

CRITICAL: Respond ONLY with valid JSON. Extract actual numbers from the document. If data is missing, use reasonable estimates based on context."""
    
    result_text = await get_llm_analysis(prompt)
    
    # Try to parse as JSON, fallback to structured text
    try:
        result = json.loads(result_text)
    except:
        result = {"analysis": result_text}
    
    # Save analysis
    analysis = Analysis(
        document_id=request.document_id,
        analysis_type="statement",
        content=request.content[:5000],  # Store first 5000 chars
        result=result
    )
    
    analysis_dict = analysis.model_dump()
    analysis_dict['created_at'] = analysis_dict['created_at'].isoformat()
    await db.analyses.insert_one(analysis_dict)
    
    logger.info(f"Analysis saved with ID: {analysis.id}")
    
    return {"success": True, "analysis_id": analysis.id, "result": result}

@api_router.post("/analyze/earnings")
async def analyze_earnings(request: AnalysisRequest):
    """Analyze earnings call transcript"""
    prompt = f"""Analyze this earnings call transcript:

{request.content}

Provide:
1. Quarterly performance summary
2. Management tone and sentiment (positive/neutral/negative)
3. Forward guidance analysis
4. Key risks mentioned
5. Key opportunities mentioned
6. What investors need to know

Format as JSON with: summary, sentiment, guidance, risks, opportunities, investor_takeaways"""
    
    result_text = await get_llm_analysis(prompt)
    
    try:
        result = json.loads(result_text)
    except:
        result = {"analysis": result_text}
    
    analysis = Analysis(
        document_id=request.document_id,
        analysis_type="earnings",
        content=request.content,
        result=result
    )
    
    analysis_dict = analysis.model_dump()
    analysis_dict['created_at'] = analysis_dict['created_at'].isoformat()
    await db.analyses.insert_one(analysis_dict)
    
    return {"success": True, "analysis_id": analysis.id, "result": result}

@api_router.post("/analyze/industry")
async def analyze_industry(request: AnalysisRequest):
    """Analyze industry dynamics"""
    prompt = f"""Analyze this industry: {request.content}

Provide comprehensive analysis:
1. Industry overview and value chain
2. Key drivers (macro, supply/demand, costs)
3. Competitive landscape
4. Top players and market share
5. Margin trends and profit pools
6. Growth opportunities
7. Key risks and challenges

Format as JSON with: overview, value_chain, drivers, competitive_landscape, top_players, margin_trends, growth_opportunities, risks"""
    
    result_text = await get_llm_analysis(prompt)
    
    try:
        result = json.loads(result_text)
    except:
        result = {"analysis": result_text}
    
    analysis = Analysis(
        analysis_type="industry",
        content=request.content,
        result=result
    )
    
    analysis_dict = analysis.model_dump()
    analysis_dict['created_at'] = analysis_dict['created_at'].isoformat()
    await db.analyses.insert_one(analysis_dict)
    
    return {"success": True, "analysis_id": analysis.id, "result": result}

@api_router.post("/analyze/red-flags")
async def analyze_red_flags(request: AnalysisRequest):
    """Detect accounting red flags"""
    prompt = f"""Analyze this financial data for red flags:

{request.content}

Detect and explain:
1. Revenue recognition anomalies
2. Working capital issues
3. Cash flow vs earnings mismatches
4. Aggressive expense capitalization
5. Margin compression patterns
6. Poor cash quality indicators
7. Off-balance-sheet liabilities
8. Interest coverage concerns

For each red flag found, provide:
- Severity (high/medium/low)
- Description
- Impact assessment

Also provide an overall risk score (0-100, higher = more risk)

Format as JSON with: red_flags (array), overall_risk_score, summary"""
    
    result_text = await get_llm_analysis(prompt)
    
    try:
        result = json.loads(result_text)
    except:
        result = {"analysis": result_text, "overall_risk_score": 50}
    
    analysis = Analysis(
        document_id=request.document_id,
        analysis_type="red_flags",
        content=request.content,
        result=result
    )
    
    analysis_dict = analysis.model_dump()
    analysis_dict['created_at'] = analysis_dict['created_at'].isoformat()
    await db.analyses.insert_one(analysis_dict)
    
    return {"success": True, "analysis_id": analysis.id, "result": result}

@api_router.post("/analyze/kpi")
async def analyze_kpis(request: AnalysisRequest):
    """Analyze business KPIs"""
    prompt = f"""Analyze these business KPIs:

{request.content}

Identify and analyze:
1. Key performance indicators relevant to this business
2. Trends (YoY, QoQ) for each KPI
3. What drives each KPI
4. Industry benchmarks if applicable
5. Areas of strength and concern

Format as JSON with: kpis (array with name, value, trend, drivers), strengths, concerns"""
    
    result_text = await get_llm_analysis(prompt)
    
    try:
        result = json.loads(result_text)
    except:
        result = {"analysis": result_text}
    
    analysis = Analysis(
        document_id=request.document_id,
        analysis_type="kpi",
        content=request.content,
        result=result
    )
    
    analysis_dict = analysis.model_dump()
    analysis_dict['created_at'] = analysis_dict['created_at'].isoformat()
    await db.analyses.insert_one(analysis_dict)
    
    return {"success": True, "analysis_id": analysis.id, "result": result}

@api_router.post("/analyze/valuation")
async def analyze_valuation(request: AnalysisRequest):
    """Generate valuation narrative"""
    style = request.document_id or "default"  # Use document_id field to pass style
    
    style_prompts = {
        "buffett": "Analyze this investment like Warren Buffett would - focus on moat, management quality, and long-term value creation.",
        "burry": "Analyze this investment like Michael Burry would - focus on deep value, contrarian indicators, and market inefficiencies.",
        "cfa": "Analyze this investment like a CFA Level 3 candidate - comprehensive, structured, with proper valuation frameworks.",
        "default": "Provide comprehensive investment analysis."
    }
    
    style_instruction = style_prompts.get(style, style_prompts["default"])
    
    prompt = f"""{style_instruction}

{request.content}

Provide:
1. Investment thesis
2. Growth drivers analysis
3. Margin outlook
4. Competitive moat assessment
5. Management quality
6. Capital allocation track record
7. Bull case scenario
8. Bear case scenario
9. Fair value assessment (directional: undervalued/fairly valued/overvalued)
10. Key risks to the thesis

Format as JSON with these sections."""
    
    result_text = await get_llm_analysis(prompt)
    
    try:
        result = json.loads(result_text)
    except:
        result = {"analysis": result_text}
    
    analysis = Analysis(
        analysis_type="valuation",
        content=request.content,
        result=result
    )
    
    analysis_dict = analysis.model_dump()
    analysis_dict['created_at'] = analysis_dict['created_at'].isoformat()
    await db.analyses.insert_one(analysis_dict)
    
    return {"success": True, "analysis_id": analysis.id, "result": result}

@api_router.post("/analyze/consistency")
async def analyze_consistency(request: AnalysisRequest):
    """Financial Consistency Engine - Cross-statement reconciliation and forensic analysis"""
    
    prompt = f"""You are a forensic financial analyst performing cross-statement reconciliation.

FINANCIAL DATA:
{request.content[:15000]}

ANALYSIS TASKS:
1. Reconcile net income with operating cash flow
2. Identify sources of divergence (working capital, accruals, non-cash items)
3. Assess whether divergence is normal or concerning
4. Analyze balance sheet changes vs income statement
5. Detect potential earnings manipulation indicators

RESPOND IN VALID JSON FORMAT:

{{
  "summary": "Executive summary of financial statement consistency (3-4 sentences)",
  
  "reconciliation": {{
    "net_income": "Value from income statement",
    "operating_cash_flow": "Value from cash flow statement",
    "difference": "Calculated difference",
    "difference_percentage": "Percentage divergence"
  }},
  
  "divergence_drivers": [
    {{
      "factor": "Working Capital Changes",
      "impact": "Dollar amount impact",
      "direction": "Increase/Decrease",
      "explanation": "Why this creates divergence",
      "normal": true/false
    }},
    {{
      "factor": "Accounts Receivable Growth",
      "impact": "Dollar amount",
      "direction": "Increase",
      "explanation": "AR growing faster than revenue",
      "normal": false
    }},
    {{
      "factor": "Depreciation & Amortization",
      "impact": "Dollar amount",
      "direction": "Add back",
      "explanation": "Non-cash expense",
      "normal": true
    }}
  ],
  
  "working_capital_analysis": {{
    "accounts_receivable": {{
      "change": "Dollar change",
      "days_sales_outstanding": "DSO metric if available",
      "concern_level": "Low/Medium/High"
    }},
    "inventory": {{
      "change": "Dollar change", 
      "days_inventory": "DIO metric if available",
      "concern_level": "Low/Medium/High"
    }},
    "accounts_payable": {{
      "change": "Dollar change",
      "days_payable": "DPO metric if available",
      "concern_level": "Low/Medium/High"
    }}
  }},
  
  "balance_sheet_pressure": {{
    "debt_increase": "Change in total debt",
    "equity_changes": "Changes in shareholders equity",
    "asset_quality": "Assessment of asset composition",
    "concerns": ["List of specific balance sheet concerns"]
  }},
  
  "manipulation_indicators": [
    {{
      "indicator": "Revenue Recognition Timing",
      "evidence": "Specific evidence found",
      "severity": "Low/Medium/High"
    }},
    {{
      "indicator": "Expense Capitalization",
      "evidence": "Pattern observed",
      "severity": "Low/Medium/High"
    }}
  ],
  
  "cash_quality_score": 85,
  
  "risk_assessment": {{
    "overall_risk": "Low/Medium/High",
    "primary_concerns": ["List top 3 concerns"],
    "mitigating_factors": ["Positive factors"],
    "recommendation": "Investor action recommendation"
  }},
  
  "confidence_score": 82,
  
  "key_findings": [
    "Most important finding 1",
    "Most important finding 2", 
    "Most important finding 3"
  ],
  
  "reconciliation_chart_data": [
    {{"item": "Net Income", "value": 320}},
    {{"item": "Add: Depreciation", "value": 50}},
    {{"item": "Less: WC Increase", "value": -80}},
    {{"item": "Other Adjustments", "value": 10}},
    {{"item": "Operating Cash Flow", "value": 300}}
  ]
}}

CRITICAL RULES:
- Use ONLY provided data, no assumptions
- Quantify ALL differences in dollars
- Clearly state if data is missing
- Assign confidence based on data completeness
- Be specific about what creates concern vs what's normal"""
    
    result_text = await get_llm_analysis(prompt)
    
    try:
        result = json.loads(result_text)
    except:
        result = {"analysis": result_text}
    
    # Save analysis
    analysis = Analysis(
        document_id=request.document_id,
        analysis_type="consistency",
        content=request.content[:5000],
        result=result
    )
    
    analysis_dict = analysis.model_dump()
    analysis_dict['created_at'] = analysis_dict['created_at'].isoformat()
    await db.analyses.insert_one(analysis_dict)
    
    logger.info(f"Consistency analysis saved with ID: {analysis.id}")
    
    return {"success": True, "analysis_id": analysis.id, "result": result}

@api_router.post("/analyze/earnings-quality")
async def analyze_earnings_quality(request: AnalysisRequest):
    """Earnings Quality Score - Distinguishes real vs engineered earnings"""
    
    prompt = f"""You are an expert at evaluating corporate earnings quality.

FINANCIAL DATA:
{request.content[:15000]}

ANALYSIS TASKS:
1. Compare earnings growth to cash flow growth
2. Analyze margin expansion sources and sustainability
3. Identify aggressive accounting indicators
4. Detect revenue recognition timing issues
5. Assess quality of reported EPS

SCORING CRITERIA:
- High quality (8-10): Strong cash conversion, sustainable margins, conservative accounting
- Medium quality (5-7): Moderate concerns, some engineering, mixed signals
- Low quality (1-4): Significant red flags, aggressive accounting, earnings trap risk

RESPOND IN VALID JSON FORMAT:

{{
  "earnings_quality_score": 7,
  
  "summary": "Executive summary of earnings quality assessment (3-4 sentences)",
  
  "score_justification": "Detailed explanation of why this specific score was assigned",
  
  "eps_vs_cash_analysis": {{
    "eps_growth": "X% growth YoY",
    "operating_cash_flow_growth": "Y% growth YoY",
    "divergence": "Z percentage points",
    "assessment": "Growing divergence signals quality concerns / Strong alignment indicates real earnings",
    "concern_level": "Low/Medium/High"
  }},
  
  "margin_analysis": {{
    "gross_margin_trend": "Expanding/Stable/Contracting",
    "operating_margin_trend": "Expanding/Stable/Contracting",
    "margin_expansion_drivers": [
      "Operational efficiency improvements",
      "Pricing power",
      "Cost reduction initiatives"
    ],
    "sustainability_assessment": "Margins appear sustainable / Margins may face pressure",
    "red_flags": ["One-time benefits inflating margins", "Unsustainable cost cuts"]
  }},
  
  "revenue_quality": {{
    "revenue_recognition_assessment": "Conservative / Aggressive / Normal",
    "accounts_receivable_quality": {{
      "dso_trend": "Improving/Stable/Deteriorating",
      "dso_value": "X days",
      "concern": "AR growing faster than revenue suggests pull-forward"
    }},
    "revenue_concentration_risk": "Low/Medium/High",
    "one_time_items": ["List any non-recurring revenue items"]
  }},
  
  "aggressive_accounting_indicators": [
    {{
      "indicator": "Revenue Recognition Timing",
      "evidence": "Specific pattern observed",
      "severity": "Low/Medium/High",
      "impact_on_score": -1
    }},
    {{
      "indicator": "Capitalized Expenses",
      "evidence": "Rising capex-to-maintenance ratio",
      "severity": "Medium",
      "impact_on_score": -0.5
    }},
    {{
      "indicator": "Inventory Buildup",
      "evidence": "Inventory growing faster than COGS",
      "severity": "Medium",
      "impact_on_score": -0.5
    }}
  ],
  
  "key_red_flags": [
    "EPS growing 20% while cash flow flat - major quality concern",
    "Margins boosted by one-time restructuring benefits",
    "DSO increased from 45 to 58 days - aggressive collections"
  ],
  
  "positive_signals": [
    "Cash conversion consistently above 100%",
    "Conservative revenue recognition policies",
    "Sustainable margin expansion from pricing power"
  ],
  
  "earnings_trap_risk": {{
    "risk_level": "Low/Medium/High",
    "probability": "Percentage likelihood of earnings disappointment",
    "key_vulnerabilities": ["What could cause earnings to disappoint"],
    "warning_signs": ["Early indicators to watch"]
  }},
  
  "quality_breakdown": {{
    "cash_conversion_quality": 8,
    "margin_quality": 7,
    "revenue_quality": 6,
    "balance_sheet_quality": 8,
    "accounting_conservatism": 7
  }},
  
  "comparison_chart_data": [
    {{"metric": "EPS Growth", "value": 15, "benchmark": 10}},
    {{"metric": "CFO Growth", "value": 12, "benchmark": 10}},
    {{"metric": "Gross Margin", "value": 48, "benchmark": 45}},
    {{"metric": "Op Margin", "value": 18, "benchmark": 15}},
    {{"metric": "Cash Conversion", "value": 105, "benchmark": 100}}
  ]},
  
  "verdict": "High Quality / Medium Quality / Low Quality - Detailed conclusion",
  
  "investor_action": "Buy/Hold/Avoid with specific reasoning",
  
  "confidence_level": 85
}}

CRITICAL RULES:
- Score from 1-10 only (integers)
- Quantify all growth rates and trends
- Be specific about red flags with evidence
- Distinguish temporary vs structural issues
- Consider industry context where applicable"""
    
    result_text = await get_llm_analysis(prompt)
    
    try:
        result = json.loads(result_text)
    except:
        result = {"analysis": result_text}
    
    # Save analysis
    analysis = Analysis(
        document_id=request.document_id,
        analysis_type="earnings_quality",
        content=request.content[:5000],
        result=result
    )
    
    analysis_dict = analysis.model_dump()
    analysis_dict['created_at'] = analysis_dict['created_at'].isoformat()
    await db.analyses.insert_one(analysis_dict)
    
    logger.info(f"Earnings quality analysis saved with ID: {analysis.id}")
    
    return {"success": True, "analysis_id": analysis.id, "result": result}

@api_router.post("/analyze/income-statement")
async def analyze_income_statement(request: AnalysisRequest):
    """Specialized Income Statement Analysis"""
    
    prompt = f"""You are analyzing an Income Statement (Profit & Loss Statement).

INCOME STATEMENT DATA:
{request.content[:15000]}

Perform comprehensive P&L analysis and respond in VALID JSON:

{{
  "document_type": "Income Statement",
  "period_covered": "Quarter/Year identified",
  
  "summary": "Executive summary of income statement performance (4-5 sentences)",
  
  "revenue_analysis": {{
    "total_revenue": "Dollar amount",
    "revenue_growth_yoy": "X% growth",
    "revenue_growth_qoq": "X% growth if quarterly",
    "revenue_breakdown": [
      {{"segment": "Product Revenue", "amount": "$X", "percentage": "Y%"}},
      {{"segment": "Service Revenue", "amount": "$X", "percentage": "Y%"}}
    ],
    "revenue_quality": "Assessment of revenue sustainability and quality"
  }},
  
  "profitability_metrics": {{
    "gross_profit": "$X",
    "gross_margin": "X%",
    "gross_margin_trend": "Expanding/Stable/Contracting",
    "operating_income": "$X",
    "operating_margin": "X%",
    "operating_margin_trend": "Expanding/Stable/Contracting",
    "net_income": "$X",
    "net_margin": "X%",
    "net_margin_trend": "Expanding/Stable/Contracting",
    "ebitda": "$X (if calculable)",
    "ebitda_margin": "X%"
  }},
  
  "cost_structure": {{
    "cost_of_revenue": "$X",
    "cogs_percentage": "X% of revenue",
    "operating_expenses": "$X",
    "opex_percentage": "X% of revenue",
    "rd_expense": "$X",
    "sales_marketing": "$X",
    "general_admin": "$X",
    "cost_efficiency": "Assessment of cost management"
  }},
  
  "earnings_per_share": {{
    "basic_eps": "$X",
    "diluted_eps": "$X",
    "eps_growth": "X% vs prior period",
    "shares_outstanding": "X million",
    "share_count_trend": "Increasing/Decreasing/Stable"
  }},
  
  "margin_waterfall": [
    {{"item": "Revenue", "value": 100, "color": "#10b981"}},
    {{"item": "Less: COGS", "value": -52, "color": "#ef4444"}},
    {{"item": "Gross Profit", "value": 48, "color": "#3b82f6"}},
    {{"item": "Less: OpEx", "value": -30, "color": "#ef4444"}},
    {{"item": "Operating Income", "value": 18, "color": "#8b5cf6"}},
    {{"item": "Less: Taxes/Int", "value": -5, "color": "#ef4444"}},
    {{"item": "Net Income", "value": 13, "color": "#10b981"}}
  ],
  
  "key_trends": [
    "Most significant trend 1",
    "Most significant trend 2",
    "Most significant trend 3"
  ],
  
  "red_flags": [
    "Any concerning patterns or anomalies"
  ],
  
  "strengths": [
    "Positive aspects of performance"
  ],
  
  "yoy_comparison": {{
    "revenue_change": "+15%",
    "gross_margin_change": "+2.3 pp",
    "operating_margin_change": "+1.8 pp",
    "net_income_change": "+22%"
  }},
  
  "profitability_score": 85,
  
  "recommendation": "Investment perspective based on P&L performance"
}}

Extract actual numbers from the document. If quarterly, include QoQ comparisons. Focus on profitability drivers and margin trends."""
    
    result_text = await get_llm_analysis(prompt)
    
    try:
        result = json.loads(result_text)
    except:
        result = {"analysis": result_text}
    
    analysis = Analysis(
        document_id=request.document_id,
        analysis_type="income_statement",
        content=request.content[:5000],
        result=result
    )
    
    analysis_dict = analysis.model_dump()
    analysis_dict['created_at'] = analysis_dict['created_at'].isoformat()
    await db.analyses.insert_one(analysis_dict)
    
    return {"success": True, "analysis_id": analysis.id, "result": result}

@api_router.post("/analyze/balance-sheet")
async def analyze_balance_sheet(request: AnalysisRequest):
    """Specialized Balance Sheet Analysis"""
    
    prompt = f"""You are analyzing a Balance Sheet (Statement of Financial Position).

BALANCE SHEET DATA:
{request.content[:15000]}

Perform comprehensive balance sheet analysis and respond in VALID JSON:

{{
  "document_type": "Balance Sheet",
  "as_of_date": "Date of balance sheet",
  
  "summary": "Executive summary of financial position and health (4-5 sentences)",
  
  "assets": {{
    "total_assets": "$X",
    "current_assets": "$X",
    "current_assets_breakdown": [
      {{"item": "Cash & Equivalents", "amount": "$X"}},
      {{"item": "Accounts Receivable", "amount": "$X"}},
      {{"item": "Inventory", "amount": "$X"}},
      {{"item": "Other Current Assets", "amount": "$X"}}
    ],
    "non_current_assets": "$X",
    "non_current_breakdown": [
      {{"item": "Property, Plant & Equipment", "amount": "$X"}},
      {{"item": "Intangible Assets", "amount": "$X"}},
      {{"item": "Goodwill", "amount": "$X"}},
      {{"item": "Other Long-term Assets", "amount": "$X"}}
    ],
    "asset_quality": "Assessment of asset composition and quality"
  }},
  
  "liabilities": {{
    "total_liabilities": "$X",
    "current_liabilities": "$X",
    "current_liabilities_breakdown": [
      {{"item": "Accounts Payable", "amount": "$X"}},
      {{"item": "Short-term Debt", "amount": "$X"}},
      {{"item": "Accrued Expenses", "amount": "$X"}},
      {{"item": "Other Current Liabilities", "amount": "$X"}}
    ],
    "non_current_liabilities": "$X",
    "long_term_debt": "$X",
    "debt_structure": "Assessment of debt maturity and terms"
  }},
  
  "equity": {{
    "total_equity": "$X",
    "common_stock": "$X",
    "retained_earnings": "$X",
    "equity_structure": "Assessment of equity composition"
  }},
  
  "liquidity_ratios": {{
    "current_ratio": "X.XX",
    "quick_ratio": "X.XX",
    "cash_ratio": "X.XX",
    "working_capital": "$X",
    "liquidity_assessment": "Strong/Adequate/Weak"
  }},
  
  "leverage_ratios": {{
    "debt_to_equity": "X.XX",
    "debt_to_assets": "X.XX",
    "equity_ratio": "X.XX",
    "leverage_assessment": "Conservative/Moderate/Aggressive"
  }},
  
  "efficiency_metrics": {{
    "asset_turnover": "X.XX (if revenue available)",
    "receivables_turnover": "X.XX (if revenue available)",
    "inventory_turnover": "X.XX (if COGS available)"
  }},
  
  "balance_sheet_composition": [
    {{"category": "Current Assets", "value": 45, "color": "#10b981"}},
    {{"category": "Non-Current Assets", "value": 55, "color": "#3b82f6"}},
    {{"category": "Current Liabilities", "value": 25, "color": "#ef4444"}},
    {{"category": "Long-term Liabilities", "value": 30, "color": "#f59e0b"}},
    {{"category": "Equity", "value": 45, "color": "#8b5cf6"}}
  ],
  
  "financial_health_indicators": [
    {{"metric": "Liquidity", "score": 85, "status": "Strong"}},
    {{"metric": "Solvency", "score": 78, "status": "Good"}},
    {{"metric": "Asset Quality", "score": 82, "status": "Strong"}},
    {{"metric": "Capital Structure", "score": 75, "status": "Good"}}
  ],
  
  "key_observations": [
    "Most important observation 1",
    "Most important observation 2",
    "Most important observation 3"
  ],
  
  "red_flags": [
    "Any concerning balance sheet items or ratios"
  ],
  
  "strengths": [
    "Strong aspects of financial position"
  ],
  
  "financial_health_score": 82,
  
  "recommendation": "Assessment of balance sheet strength and risks"
}}

Extract actual numbers and calculate all ratios. Focus on liquidity, leverage, and asset quality."""
    
    result_text = await get_llm_analysis(prompt)
    
    try:
        result = json.loads(result_text)
    except:
        result = {"analysis": result_text}
    
    analysis = Analysis(
        document_id=request.document_id,
        analysis_type="balance_sheet",
        content=request.content[:5000],
        result=result
    )
    
    analysis_dict = analysis.model_dump()
    analysis_dict['created_at'] = analysis_dict['created_at'].isoformat()
    await db.analyses.insert_one(analysis_dict)
    
    return {"success": True, "analysis_id": analysis.id, "result": result}

@api_router.post("/analyze/cash-flow")
async def analyze_cash_flow(request: AnalysisRequest):
    """Specialized Cash Flow Statement Analysis"""
    
    prompt = f"""You are analyzing a Cash Flow Statement.

CASH FLOW STATEMENT DATA:
{request.content[:15000]}

Perform comprehensive cash flow analysis and respond in VALID JSON:

{{
  "document_type": "Cash Flow Statement",
  "period_covered": "Quarter/Year identified",
  
  "summary": "Executive summary of cash generation and usage (4-5 sentences)",
  
  "operating_activities": {{
    "net_cash_from_operations": "$X",
    "net_income_reconciliation": "$X starting point",
    "major_adjustments": [
      {{"item": "Depreciation & Amortization", "amount": "$X"}},
      {{"item": "Stock-based Compensation", "amount": "$X"}},
      {{"item": "Changes in Working Capital", "amount": "$X"}}
    ],
    "operating_cash_quality": "Assessment of operating cash flow quality",
    "cash_conversion_rate": "X% (OCF / Net Income)"
  }},
  
  "investing_activities": {{
    "net_cash_from_investing": "$X",
    "capital_expenditures": "$X",
    "acquisitions": "$X",
    "asset_sales": "$X",
    "capex_intensity": "X% of revenue if available",
    "investment_assessment": "Growth/Maintenance/Divestiture focused"
  }},
  
  "financing_activities": {{
    "net_cash_from_financing": "$X",
    "debt_issuance_repayment": "$X",
    "equity_issuance_buyback": "$X",
    "dividends_paid": "$X",
    "financing_strategy": "Assessment of capital allocation"
  }},
  
  "cash_position": {{
    "beginning_cash": "$X",
    "net_change_in_cash": "$X",
    "ending_cash": "$X",
    "cash_adequacy": "Strong/Adequate/Weak"
  }},
  
  "free_cash_flow": {{
    "fcf": "$X (OCF - CapEx)",
    "fcf_margin": "X% if revenue available",
    "fcf_growth": "X% vs prior period",
    "fcf_quality": "Assessment of free cash flow sustainability"
  }},
  
  "cash_flow_breakdown": [
    {{"activity": "Operating", "amount": 300, "color": "#10b981"}},
    {{"activity": "Investing", "amount": -120, "color": "#ef4444"}},
    {{"activity": "Financing", "amount": -80, "color": "#f59e0b"}},
    {{"activity": "Net Change", "amount": 100, "color": "#3b82f6"}}
  ],
  
  "working_capital_changes": [
    {{"item": "Accounts Receivable", "impact": "$X", "direction": "Use/Source"}},
    {{"item": "Inventory", "impact": "$X", "direction": "Use/Source"}},
    {{"item": "Accounts Payable", "impact": "$X", "direction": "Use/Source"}}
  ],
  
  "cash_generation_metrics": {{
    "operating_cash_margin": "X%",
    "capex_to_ocf_ratio": "X%",
    "fcf_conversion": "X%",
    "cash_return_on_assets": "X%"
  }},
  
  "key_insights": [
    "Most important cash flow insight 1",
    "Most important cash flow insight 2",
    "Most important cash flow insight 3"
  ],
  
  "concerns": [
    "Any concerning cash flow patterns"
  ],
  
  "strengths": [
    "Strong aspects of cash generation"
  ],
  
  "cash_generation_score": 88,
  
  "sustainability_assessment": "Assessment of cash flow sustainability and quality"
}}

Calculate free cash flow and all metrics. Focus on cash generation quality and sustainability."""
    
    result_text = await get_llm_analysis(prompt)
    
    try:
        result = json.loads(result_text)
    except:
        result = {"analysis": result_text}
    
    analysis = Analysis(
        document_id=request.document_id,
        analysis_type="cash_flow",
        content=request.content[:5000],
        result=result
    )
    
    analysis_dict = analysis.model_dump()
    analysis_dict['created_at'] = analysis_dict['created_at'].isoformat()
    await db.analyses.insert_one(analysis_dict)
    
    return {"success": True, "analysis_id": analysis.id, "result": result}

@api_router.post("/simulate")
async def simulate_business(request: SimulationRequest):
    """Simulate business model with different parameters"""
    prompt = f"""Run a business model simulation for {request.company_name} with these parameters:

- Revenue Growth: {request.revenue_growth}%
- Gross Margin: {request.gross_margin}%
- Operating Leverage: {request.operating_leverage}
- CapEx Intensity: {request.capex_intensity}%

Provide analysis of:
1. Impact on profitability
2. Impact on cash flows
3. Impact on valuation
4. Key risks in this scenario
5. Sensitivity analysis (what if one parameter changes by ±20%)

Format as JSON with: profitability_impact, cash_flow_impact, valuation_impact, risks, sensitivity_analysis"""
    
    result_text = await get_llm_analysis(prompt)
    
    try:
        result = json.loads(result_text)
    except:
        result = {"analysis": result_text}
    
    return {"success": True, "result": result}

@api_router.post("/portfolio/analyze")
async def analyze_portfolio(portfolio: PortfolioCreate):
    """Analyze portfolio composition and risk"""
    holdings_text = json.dumps(portfolio.holdings, indent=2)
    
    prompt = f"""Analyze this investment portfolio:

{holdings_text}

Provide:
1. Diversification score (0-100)
2. Sector breakdown
3. Risk concentration analysis
4. Correlation risk
5. Factor exposure (growth, value, quality, etc.)
6. Recession sensitivity
7. Recommendations for improvement
8. Overall portfolio health score (0-100)

Format as JSON with these sections."""
    
    result_text = await get_llm_analysis(prompt)
    
    try:
        result = json.loads(result_text)
    except:
        result = {"analysis": result_text}
    
    # Save portfolio
    portfolio_obj = Portfolio(
        name=portfolio.name,
        holdings=portfolio.holdings
    )
    
    portfolio_dict = portfolio_obj.model_dump()
    portfolio_dict['created_at'] = portfolio_dict['created_at'].isoformat()
    await db.portfolios.insert_one(portfolio_dict)
    
    return {"success": True, "portfolio_id": portfolio_obj.id, "analysis": result}

@api_router.get("/documents")
async def get_documents():
    """Get all uploaded documents"""
    docs = await db.documents.find({}, {"_id": 0, "content": 0}).to_list(100)
    return {"documents": docs}

@api_router.get("/documents/{document_id}")
async def get_document(document_id: str):
    """Get a specific document with full content"""
    doc = await db.documents.find_one({"id": document_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc

@api_router.get("/analyses")
async def get_analyses():
    """Get all analyses"""
    analyses = await db.analyses.find({}, {"_id": 0}).to_list(100)
    return {"analyses": analyses}

@api_router.get("/portfolios")
async def get_portfolios():
    """Get all portfolios"""
    portfolios = await db.portfolios.find({}, {"_id": 0}).to_list(100)
    return {"portfolios": portfolios}

@api_router.post("/companies")
async def create_company(company: CompanyCreate):
    """Create/track a company"""
    company_obj = Company(**company.model_dump(), tracked=True)
    
    company_dict = company_obj.model_dump()
    company_dict['created_at'] = company_dict['created_at'].isoformat()
    await db.companies.insert_one(company_dict)
    
    return {"success": True, "company_id": company_obj.id}

@api_router.get("/companies")
async def get_companies():
    """Get all tracked companies"""
    companies = await db.companies.find({"tracked": True}, {"_id": 0}).to_list(100)
    return {"companies": companies}

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()